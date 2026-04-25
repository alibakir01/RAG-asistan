"""
Endüstri Mühendisliği belgeleri için ingest modülü.
Çıktı: data/processed/chunks_endustri.jsonl
"""
from __future__ import annotations

import csv
import json
import re
from dataclasses import asdict, dataclass, field
from pathlib import Path
from typing import Iterator

from docx import Document
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
RAW = ROOT / "data" / "raw"
OUT = ROOT / "data" / "processed"
OUT.mkdir(parents=True, exist_ok=True)

BOLUM = "endustri"


@dataclass
class Chunk:
    id: str
    text: str
    metadata: dict = field(default_factory=dict)


# ----------------------------- Yardımcılar -----------------------------

COL_ALIASES = {
    "ders_kodu": ["ders kodu", "kod", "code"],
    "ders_adi": ["ders adı", "ders adi", "course name", "course", "ad", "name"],
    "on_sart": ["ön şart", "on sart", "ön sart", "prereq", "prerequisite", "prerequisites"],
    "teorik": ["teorik", "teo.", "teo", "theoretical", "theory", "lec.", "lec", "lecture", "t"],
    "lab": ["laboratuvar", "lab", "practical", "p"],
    "kredi": ["kredi", "credit", "credits"],
    "akts": ["akts", "ects"],
}


def _col_map(header_cells: list[str]) -> dict:
    norm = {}
    for i, h in enumerate(header_cells):
        hl = h.strip().lower()
        for key, alts in COL_ALIASES.items():
            if any(hl == a or hl.startswith(a + " ") or hl == a.replace(".", "") for a in alts):
                norm.setdefault(key, i)
                break
    return norm


def _cell(cells: list[str], cmap: dict, key: str) -> str:
    i = cmap.get(key)
    if i is None or i >= len(cells):
        return ""
    return cells[i].strip()


def _blankish(v: str) -> bool:
    """Boş veya tire benzeri değer mi?"""
    return v is None or str(v).strip() in ("", "-", "—", "–", "N/A", "n/a")


def _synth_code(ders_adi: str, kind: str = "sec") -> str:
    """Kaynakta kod olmayan elective satırlarına sentetik kod üret."""
    ad = (ders_adi or "").lower()
    if "non-technical" in ad or "non technical" in ad:
        return "IE-ELEC-NT"
    if "technical" in ad:
        return "IE-ELEC-TEK"
    if "senior" in ad:
        return "IE-ELEC-SR"
    if "math" in ad:
        return "IE-ELEC-MATH"
    return "IE-ELEC"


def _fmt_hours(teorik: str, lab: str, kredi: str, akts: str, kod: str = "", ders_adi: str = "") -> str:
    """Ders saat/kredi bilgilerini okunaklı Türkçe metne çevir.
    Boş değerler 'belirlenmemiş' olarak, '-' değerler ders tipine göre ('stajlı ders'
    veya 'seçmeli havuzdan atanır') olarak yazılır."""
    t_blank = _blankish(teorik)
    l_blank = _blankish(lab)
    k_blank = _blankish(kredi)
    a_blank = _blankish(akts)

    ad_low = (ders_adi or "").lower()
    kod_up = (kod or "").upper().replace(" ", "")
    is_internship = (
        "summer training" in ad_low or "summer practice" in ad_low
        or kod_up in ("IE197", "IE297", "IE397", "IE298", "IE398", "IE400")
        or "workplace experience" in ad_low
    )
    is_elective_placeholder = (
        not kod or kod.startswith("IE-ELEC")
        or kod_up in ("IETX", "IENX", "IESX", "IEDX", "IEMX")
        or "elective" in ad_low
    )

    def _val(v: str, is_blank: bool) -> str:
        if not is_blank:
            return str(v).strip()
        if is_internship:
            return "—"  # staj dersleri için sınıf saati yok
        if is_elective_placeholder:
            return "seçim sonrası belirlenir"
        return "belirtilmemiş"

    t = _val(teorik, t_blank)
    l = _val(lab, l_blank)
    k = _val(kredi, k_blank)
    a = _val(akts, a_blank)

    parts = [f"Teorik: {t}", f"Lab: {l}", f"Kredi: {k}", f"AKTS: {a}"]
    base = ", ".join(parts) + "."
    if is_internship and (t_blank or l_blank or k_blank):
        base += " (Staj dersi — sınıf içi ders saati yoktur.)"
    elif is_elective_placeholder and (t_blank or k_blank):
        base += " (Seçmeli havuzdan seçilen dersin saat/kredi bilgileri o derse göre değişir.)"
    return base


def _iter_body(doc):
    p_idx, t_idx = 0, 0
    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:p"):
            yield "p", doc.paragraphs[p_idx]
            p_idx += 1
        elif child.tag == qn("w:tbl"):
            yield "t", doc.tables[t_idx]
            t_idx += 1


# ----------------------------- Müfredat: 16-21 ve 21-25 (EN paragraflı) -----------------------------

SEM_EN_RE = re.compile(
    r"(FIRST|SECOND|THIRD|FOURTH)\s*YEAR\s*/\s*"
    r"(FIRST|SECOND|THIRD|FOURTH|FIFTH|SIXTH|SEVENTH|EIGHTH|EIGTH)\s*SEMESTER",
    re.IGNORECASE,
)
YEAR_MAP = {"FIRST": 1, "SECOND": 2, "THIRD": 3, "FOURTH": 4}
SEM_MAP = {
    "FIRST": 1, "SECOND": 2, "THIRD": 3, "FOURTH": 4,
    "FIFTH": 5, "SIXTH": 6, "SEVENTH": 7, "EIGHTH": 8, "EIGTH": 8,
}
IE_POOL_RE = re.compile(r"(electives?|concentration|year long)", re.IGNORECASE)


def parse_ie_mufredat_legacy(path: Path, mufredat_yili: str) -> Iterator[Chunk]:
    doc = Document(path)
    cur_donem = cur_yil = cur_sezon = None
    cur_pool = None
    semester_courses: dict[int, list[str]] = {}
    semester_meta: dict[int, dict] = {}

    for kind, el in _iter_body(doc):
        if kind == "p":
            txt = el.text.strip()
            if not txt:
                continue
            m = SEM_EN_RE.search(txt)
            if m:
                yil = YEAR_MAP[m.group(1).upper()]
                donem = SEM_MAP[m.group(2).upper()]
                cur_yil = yil
                cur_donem = donem
                cur_sezon = "Güz" if donem % 2 == 1 else "Bahar"
                cur_pool = None
            elif IE_POOL_RE.search(txt) and len(txt) < 100 and "must be taken" not in txt.lower():
                cur_pool = txt
                cur_donem = None  # havuz geldi
        elif kind == "t":
            rows = el.rows
            if not rows:
                continue
            header = [c.text.strip() for c in rows[0].cells]
            cmap = _col_map(header)
            if "ders_kodu" not in cmap or "ders_adi" not in cmap:
                continue
            for ri, row in enumerate(rows[1:], start=1):
                cells = [c.text.strip() for c in row.cells]
                if not any(cells):
                    continue
                ders_kodu = _cell(cells, cmap, "ders_kodu")
                ders_adi = _cell(cells, cmap, "ders_adi")
                if not ders_adi or ders_adi.upper() in ("TOTAL", "TOPLAM"):
                    continue
                teorik = _cell(cells, cmap, "teorik")
                lab = _cell(cells, cmap, "lab")
                kredi = _cell(cells, cmap, "kredi")
                akts = _cell(cells, cmap, "akts")
                on_sart = _cell(cells, cmap, "on_sart").strip("-").strip()

                kod_effective = ders_kodu or _synth_code(ders_adi)
                hours_str = _fmt_hours(teorik, lab, kredi, akts, kod_effective, ders_adi)

                if cur_donem:
                    text = (
                        f"{mufredat_yili} Endüstri müfredatı — {cur_yil}. yıl {cur_sezon} dönemi "
                        f"({cur_donem}. dönem): {kod_effective} {ders_adi}. "
                        f"{hours_str} "
                        f"Ön şart: {on_sart if on_sart else 'yok'}."
                    )
                    md = {
                        "tip": "mufredat",
                        "mufredat_yili": mufredat_yili,
                        "donem": cur_donem,
                        "yil": cur_yil,
                        "sezon": cur_sezon,
                        "ders_kodu": kod_effective,
                        "ders_adi": ders_adi,
                        "on_sart": on_sart,
                        "teorik": teorik,
                        "lab": lab,
                        "kredi": kredi,
                        "akts": akts,
                        "kaynak": path.name,
                    }
                    cid = f"ie_muf_{mufredat_yili}_d{cur_donem}_r{ri}_{kod_effective.replace(' ','')}"
                    short_line = (
                        f"- {kod_effective} | {ders_adi} | "
                        f"T:{teorik or '—'}, L:{lab or '—'}, Kredi:{kredi or '—'}, AKTS:{akts or '—'} | "
                        f"Ön şart: {on_sart or 'yok'}"
                    )
                    semester_courses.setdefault(cur_donem, []).append(short_line)
                    semester_meta.setdefault(cur_donem, {"yil": cur_yil, "sezon": cur_sezon})
                else:
                    pool = cur_pool or "IE Seçmeli Havuz"
                    text = (
                        f"{mufredat_yili} Endüstri müfredatı — {pool} havuzu: {kod_effective} {ders_adi}. "
                        f"{hours_str} "
                        f"Ön şart: {on_sart if on_sart else 'yok'}."
                    )
                    md = {
                        "tip": "secmeli_havuz",
                        "mufredat_yili": mufredat_yili,
                        "havuz": pool,
                        "ders_kodu": kod_effective,
                        "ders_adi": ders_adi,
                        "on_sart": on_sart,
                        "teorik": teorik,
                        "lab": lab,
                        "kredi": kredi,
                        "akts": akts,
                        "kaynak": path.name,
                    }
                    cid = f"ie_muf_{mufredat_yili}_pool_{pool[:12].replace(' ','_')}_r{ri}_{kod_effective.replace(' ','')}"
                yield Chunk(id=cid, text=text, metadata=md)

    # Dönem özeti (tek chunk — LLM'in atlaması zor)
    for donem, lines in sorted(semester_courses.items()):
        meta = semester_meta[donem]
        body = (
            f"Endüstri Mühendisliği {mufredat_yili} Müfredatı — "
            f"{meta['yil']}. sınıf {meta['sezon']} yarıyılı ({donem}. dönem) — Toplam {len(lines)} ders:\n"
            + "\n".join(lines)
        )
        yield Chunk(
            id=f"ie_muf_{mufredat_yili}_sem_summary_d{donem}",
            text=body,
            metadata={
                "tip": "donem_ozet",
                "mufredat_yili": mufredat_yili,
                "donem": donem,
                "yil": meta["yil"],
                "sezon": meta["sezon"],
                "ders_sayisi": len(lines),
                "kaynak": path.name,
            },
        )


# ----------------------------- Müfredat: 25-26 (merged-cell header) -----------------------------

MERGED_SEM_RE = re.compile(
    r"(Freshman|Sophomore|Junior|Senior)\s*Year\s*/\s*(FALL|SPRING|FA|SP)",
    re.IGNORECASE,
)
YEAR_WORD_MAP = {"freshman": 1, "sophomore": 2, "junior": 3, "senior": 4}


def parse_ie_mufredat_new(path: Path, mufredat_yili: str) -> Iterator[Chunk]:
    doc = Document(path)
    pool_paragraph = None
    semester_courses: dict[int, list[str]] = {}
    semester_meta: dict[int, dict] = {}

    for kind, el in _iter_body(doc):
        if kind == "p":
            txt = el.text.strip()
            if not txt:
                continue
            # Havuz işaretçileri (IEDX-, IEMX-, IESX-, GLB Electives, Technical Electives offered by IE, ...)
            if re.search(r"electives?", txt, re.IGNORECASE) and len(txt) < 120:
                pool_paragraph = txt
            elif not MERGED_SEM_RE.search(txt):
                # havuz state'i bozma
                pass
        elif kind == "t":
            rows = el.rows
            if not rows:
                continue

            # 25-26 style: ilk satır merged cell (Freshman Year / FA ...), 2. satır gerçek header
            # Legacy style: 1. satır header zaten
            hdr1 = [c.text.strip() for c in rows[0].cells]
            hdr2 = [c.text.strip() for c in rows[1].cells] if len(rows) > 1 else []

            sem_match = MERGED_SEM_RE.search(hdr1[0] if hdr1 else "")
            if sem_match and len(rows) >= 2:
                # Semester tablosu
                cur_yil = YEAR_WORD_MAP[sem_match.group(1).lower()]
                fa_sp = sem_match.group(2).upper()
                cur_sezon = "Güz" if fa_sp.startswith("FA") else "Bahar"
                cur_donem = (cur_yil - 1) * 2 + (1 if cur_sezon == "Güz" else 2)
                data_rows = rows[2:]
                cmap = _col_map(hdr2)
                is_pool = False
            else:
                # Havuz tablosu — ilk satır header
                cur_yil = cur_donem = None
                cur_sezon = None
                data_rows = rows[1:]
                cmap = _col_map(hdr1)
                is_pool = True

            if "ders_kodu" not in cmap or "ders_adi" not in cmap:
                continue

            for ri, row in enumerate(data_rows, start=1):
                cells = [c.text.strip() for c in row.cells]
                if not any(cells):
                    continue
                ders_kodu = _cell(cells, cmap, "ders_kodu")
                ders_adi = _cell(cells, cmap, "ders_adi")
                if not ders_adi or ders_adi.upper() in ("TOTAL", "TOPLAM"):
                    continue
                teorik = _cell(cells, cmap, "teorik")
                lab = _cell(cells, cmap, "lab")
                kredi = _cell(cells, cmap, "kredi")
                akts = _cell(cells, cmap, "akts")
                on_sart = _cell(cells, cmap, "on_sart").strip("-").strip()

                kod_effective = ders_kodu or _synth_code(ders_adi)
                hours_str = _fmt_hours(teorik, lab, kredi, akts, kod_effective, ders_adi)

                if is_pool:
                    pool = pool_paragraph or "IE Seçmeli Havuz"
                    text = (
                        f"{mufredat_yili} Endüstri müfredatı — {pool} havuzu: "
                        f"{kod_effective} {ders_adi}. "
                        f"{hours_str} "
                        f"Ön şart: {on_sart if on_sart else 'yok'}."
                    )
                    md = {
                        "tip": "secmeli_havuz",
                        "mufredat_yili": mufredat_yili,
                        "havuz": pool,
                        "ders_kodu": kod_effective,
                        "ders_adi": ders_adi,
                        "on_sart": on_sart,
                        "teorik": teorik,
                        "lab": lab,
                        "kredi": kredi,
                        "akts": akts,
                        "kaynak": path.name,
                    }
                    cid = f"ie_muf_{mufredat_yili}_pool_{pool[:12].replace(' ','_')}_r{ri}_{kod_effective.replace(' ','')}"
                else:
                    text = (
                        f"{mufredat_yili} Endüstri müfredatı — {cur_yil}. yıl {cur_sezon} dönemi "
                        f"({cur_donem}. dönem): {kod_effective} {ders_adi}. "
                        f"{hours_str} "
                        f"Ön şart: {on_sart if on_sart else 'yok'}."
                    )
                    md = {
                        "tip": "mufredat",
                        "mufredat_yili": mufredat_yili,
                        "donem": cur_donem,
                        "yil": cur_yil,
                        "sezon": cur_sezon,
                        "ders_kodu": kod_effective,
                        "ders_adi": ders_adi,
                        "on_sart": on_sart,
                        "teorik": teorik,
                        "lab": lab,
                        "kredi": kredi,
                        "akts": akts,
                        "kaynak": path.name,
                    }
                    cid = f"ie_muf_{mufredat_yili}_d{cur_donem}_r{ri}_{kod_effective.replace(' ','')}"
                    short_line = (
                        f"- {kod_effective} | {ders_adi} | "
                        f"T:{teorik or '—'}, L:{lab or '—'}, Kredi:{kredi or '—'}, AKTS:{akts or '—'} | "
                        f"Ön şart: {on_sart or 'yok'}"
                    )
                    semester_courses.setdefault(cur_donem, []).append(short_line)
                    semester_meta.setdefault(cur_donem, {"yil": cur_yil, "sezon": cur_sezon})

                yield Chunk(id=cid, text=text, metadata=md)

    # Dönem özeti (tek chunk)
    for donem, lines in sorted(semester_courses.items()):
        meta = semester_meta[donem]
        body = (
            f"Endüstri Mühendisliği {mufredat_yili} Müfredatı — "
            f"{meta['yil']}. sınıf {meta['sezon']} yarıyılı ({donem}. dönem) — Toplam {len(lines)} ders:\n"
            + "\n".join(lines)
        )
        yield Chunk(
            id=f"ie_muf_{mufredat_yili}_sem_summary_d{donem}",
            text=body,
            metadata={
                "tip": "donem_ozet",
                "mufredat_yili": mufredat_yili,
                "donem": donem,
                "yil": meta["yil"],
                "sezon": meta["sezon"],
                "ders_sayisi": len(lines),
                "kaynak": path.name,
            },
        )


# ----------------------------- IE Müfredat: Temiz CSV (2016) -----------------------------

IE_SEZON_BY_DONEM = {1: "Güz", 2: "Bahar", 3: "Güz", 4: "Bahar", 5: "Güz", 6: "Bahar", 7: "Güz", 8: "Bahar"}


def parse_ie_mufredat_csv(path: Path, mufredat_yili: str) -> Iterator[Chunk]:
    """Endüstri müfredatı için elle temizlenmiş CSV: Dönem,Ders Kodu,Ders Adı,Ön Şart,Teorik,Lab.,Kredi,AKTS"""
    semester_courses: dict[int, list[str]] = {}
    semester_meta: dict[int, dict] = {}

    with path.open(encoding="utf-8") as f:
        reader = csv.DictReader(f)
        for ri, row in enumerate(reader, start=1):
            try:
                donem = int(str(row.get("Dönem", "")).strip())
            except (ValueError, TypeError):
                continue
            ders_kodu = (row.get("Ders Kodu") or "").strip()
            ders_adi = (row.get("Ders Adı") or "").strip()
            on_sart = (row.get("Ön Şart") or "").strip().strip("-").strip()
            teorik = (row.get("Teorik") or "").strip()
            lab = (row.get("Lab.") or row.get("Lab") or "").strip()
            kredi = (row.get("Kredi") or "").strip()
            akts = (row.get("AKTS") or "").strip()

            if not ders_adi:
                continue
            if ders_kodu in ("-", ""):
                ders_kodu = _synth_code(ders_adi)

            yil = (donem - 1) // 2 + 1
            sezon = IE_SEZON_BY_DONEM.get(donem, "")
            hours_str = _fmt_hours(teorik, lab, kredi, akts, ders_kodu, ders_adi)

            text = (
                f"{mufredat_yili} Endüstri müfredatı — {yil}. yıl {sezon} dönemi "
                f"({donem}. dönem): {ders_kodu} {ders_adi}. "
                f"{hours_str} "
                f"Ön şart: {on_sart if on_sart else 'yok'}."
            )
            md = {
                "tip": "mufredat",
                "mufredat_yili": mufredat_yili,
                "donem": donem,
                "yil": yil,
                "sezon": sezon,
                "ders_kodu": ders_kodu,
                "ders_adi": ders_adi,
                "on_sart": on_sart,
                "teorik": teorik,
                "lab": lab,
                "kredi": kredi,
                "akts": akts,
                "kaynak": path.name,
            }
            yield Chunk(
                id=f"ie_muf_{mufredat_yili}_csv_d{donem}_r{ri}_{ders_kodu.replace(' ','')}",
                text=text,
                metadata=md,
            )

            short_line = (
                f"- {ders_kodu} | {ders_adi} | "
                f"T:{teorik or '—'}, L:{lab or '—'}, Kredi:{kredi or '—'}, AKTS:{akts or '—'} | "
                f"Ön şart: {on_sart or 'yok'}"
            )
            semester_courses.setdefault(donem, []).append(short_line)
            semester_meta.setdefault(donem, {"yil": yil, "sezon": sezon})

    for donem, lines in sorted(semester_courses.items()):
        meta = semester_meta[donem]
        body = (
            f"Endüstri Mühendisliği {mufredat_yili} Müfredatı — "
            f"{meta['yil']}. sınıf {meta['sezon']} yarıyılı ({donem}. dönem) — Toplam {len(lines)} ders:\n"
            + "\n".join(lines)
        )
        yield Chunk(
            id=f"ie_muf_{mufredat_yili}_csv_sem_summary_d{donem}",
            text=body,
            metadata={
                "tip": "donem_ozet",
                "mufredat_yili": mufredat_yili,
                "donem": donem,
                "yil": meta["yil"],
                "sezon": meta["sezon"],
                "ders_sayisi": len(lines),
                "kaynak": path.name,
            },
        )


def parse_ie_dual_secmeli_csv(path: Path, mufredat_yili: str, havuz1: str, havuz2: str) -> Iterator[Chunk]:
    """İki tablolu CSV: 1. tablo havuz1, boş satır(lar)dan sonra 2. tablo havuz2.
    İkinci tablo da kendi 'Ders Kodu, Ders Adı, ...' header'ı ile başlar."""
    with path.open(encoding="utf-8") as f:
        rows = list(csv.reader(f))

    # İki tabloya böl: header satırı 'Ders Kodu' ile başlayan satırları bul
    headers = [i for i, r in enumerate(rows) if r and r[0].strip().lower() in ("ders kodu", "kod", "code")]
    sections = []
    for hi, start in enumerate(headers):
        end = headers[hi + 1] if hi + 1 < len(headers) else len(rows)
        section_rows = rows[start:end]
        sections.append(section_rows)

    havuz_names = [havuz1, havuz2]
    for si, sec in enumerate(sections[:2]):
        if not sec:
            continue
        header = [c.strip() for c in sec[0]]
        # column index map
        cmap = {}
        for i, h in enumerate(header):
            hl = h.lower()
            if hl.startswith("ders kod") or hl == "code" or hl == "kod":
                cmap["kod"] = i
            elif hl.startswith("ders ad") or hl == "name":
                cmap["ad"] = i
            elif hl.startswith("teo") or hl == "t":
                cmap["t"] = i
            elif hl.startswith("lab") or hl == "p":
                cmap["l"] = i
            elif hl.startswith("kredi") or hl.startswith("credit"):
                cmap["k"] = i
            elif hl.startswith("akts") or hl.startswith("ects"):
                cmap["a"] = i

        havuz_label = havuz_names[si]
        courses = []
        for ri, row in enumerate(sec[1:], start=1):
            if not row or not any(c.strip() for c in row):
                continue
            kod = row[cmap.get("kod", 0)].strip() if cmap.get("kod") is not None else ""
            ad = row[cmap.get("ad", 1)].strip() if cmap.get("ad") is not None else ""
            t = row[cmap.get("t", 2)].strip() if cmap.get("t") is not None else ""
            l = row[cmap.get("l", 3)].strip() if cmap.get("l") is not None else ""
            k = row[cmap.get("k", 4)].strip() if cmap.get("k") is not None else ""
            a = row[cmap.get("a", 5)].strip() if cmap.get("a") is not None else ""
            if not ad or not kod:
                continue
            courses.append({"kod": kod, "ad": ad, "t": t, "l": l, "kredi": k, "akts": a})

            text = (
                f"Endüstri Mühendisliği {mufredat_yili} Müfredatı — {havuz_label}: "
                f"{kod} {ad}. Teorik: {t}, Lab: {l}, Kredi: {k}, AKTS: {a}. "
                f"Bu ders {havuz_label} havuzundan alınabilir ({mufredat_yili} müfredatı)."
            )
            yield Chunk(
                id=f"ie_secd_{mufredat_yili}_{si}_{kod.replace(' ','')}",
                text=text,
                metadata={
                    "tip": "secmeli_havuz",
                    "mufredat_yili": mufredat_yili,
                    "havuz": havuz_label,
                    "ders_kodu": kod,
                    "ders_adi": ad,
                    "teorik": t,
                    "lab": l,
                    "kredi": k,
                    "akts": a,
                    "kaynak": path.name,
                },
            )

        if courses:
            lines = [
                f"- {c['kod']} | {c['ad']} | T:{c['t']}, L:{c['l']}, Kredi:{c['kredi']}, AKTS:{c['akts']}"
                for c in courses
            ]
            body = (
                f"Endüstri Mühendisliği {mufredat_yili} Müfredatı — {havuz_label} "
                f"(toplam {len(courses)} ders):\n"
                + "\n".join(lines)
                + f"\n\nNot: Bu liste {mufredat_yili} müfredatına tabi öğrenciler için geçerlidir."
            )
            yield Chunk(
                id=f"ie_secd_havuz_ozet_{mufredat_yili}_{si}",
                text=body,
                metadata={
                    "tip": "secmeli_havuz_ozet",
                    "mufredat_yili": mufredat_yili,
                    "havuz": havuz_label,
                    "ders_sayisi": len(courses),
                    "kaynak": path.name,
                },
            )


def parse_ie_multi_secmeli_csv(path: Path, mufredat_yili: str, default_havuz: str = "Bölüm Seçmelisi") -> Iterator[Chunk]:
    """Çok bölümlü seçmeli CSV: birden fazla başlık satırı içerir (örn 'Matematik Seçmelileri').
    Her bölüm: opsiyonel label satırı + header satırı + ders satırları, boş satırlarla ayrılır."""
    with path.open(encoding="utf-8") as f:
        rows = list(csv.reader(f))

    sections = []  # [(havuz_label, header_row, data_rows)]
    i = 0
    current_label = default_havuz
    while i < len(rows):
        row = rows[i]
        if not row or not any(c.strip() for c in row):
            i += 1
            continue
        first = row[0].strip()
        if first.lower() in ("ders kodu", "kod", "code"):
            # Header — sonraki boş veya yeni label satırına kadar veriler
            header = row
            data = []
            j = i + 1
            while j < len(rows):
                rr = rows[j]
                if not rr or not any(c.strip() for c in rr):
                    break
                # Yeni bir label satırı mı?
                if rr[0].strip() and (len(rr) < 2 or not any(c.strip() for c in rr[1:])):
                    break
                data.append(rr)
                j += 1
            sections.append((current_label, header, data))
            i = j
        elif len(row) < 2 or not any(c.strip() for c in row[1:]):
            # Tek hücreli label satırı
            current_label = first
            i += 1
        else:
            i += 1

    for si, (havuz_label, header, data_rows) in enumerate(sections):
        cmap = {}
        for ci, h in enumerate([c.strip() for c in header]):
            hl = h.lower()
            if hl.startswith("ders kod") or hl in ("code", "kod"):
                cmap["kod"] = ci
            elif hl.startswith("ders ad") or hl == "name":
                cmap["ad"] = ci
            elif hl.startswith("ön şart") or hl.startswith("on sart") or hl.startswith("prereq"):
                cmap["on"] = ci
            elif hl.startswith("teo") or hl == "t":
                cmap["t"] = ci
            elif hl.startswith("lab") or hl == "p":
                cmap["l"] = ci
            elif hl.startswith("kredi") or hl.startswith("credit"):
                cmap["k"] = ci
            elif hl.startswith("akts") or hl.startswith("ects"):
                cmap["a"] = ci

        courses = []
        for ri, row in enumerate(data_rows, start=1):
            def _g(key, default=""):
                idx = cmap.get(key)
                if idx is None or idx >= len(row):
                    return default
                return row[idx].strip()
            kod = _g("kod")
            ad = _g("ad")
            on = _g("on").strip("-").strip()
            t = _g("t")
            l = _g("l")
            k = _g("k")
            a = _g("a")
            if not ad:
                continue
            if not kod:
                kod = _synth_code(ad)
            courses.append({"kod": kod, "ad": ad, "on": on, "t": t, "l": l, "kredi": k, "akts": a})

            text = (
                f"Endüstri Mühendisliği {mufredat_yili} Müfredatı — {havuz_label}: "
                f"{kod} {ad}. Teorik: {t or '—'}, Lab: {l or '—'}, Kredi: {k or '—'}, AKTS: {a or '—'}. "
                f"Ön şart: {on or 'yok'}. "
                f"Bu ders {havuz_label} havuzundan alınabilir ({mufredat_yili} müfredatı)."
            )
            yield Chunk(
                id=f"ie_msec_{mufredat_yili}_{si}_r{ri}_{kod.replace(' ','')}",
                text=text,
                metadata={
                    "tip": "secmeli_havuz",
                    "mufredat_yili": mufredat_yili,
                    "havuz": havuz_label,
                    "ders_kodu": kod,
                    "ders_adi": ad,
                    "on_sart": on,
                    "teorik": t,
                    "lab": l,
                    "kredi": k,
                    "akts": a,
                    "kaynak": path.name,
                },
            )

        if courses:
            lines = [
                f"- {c['kod']} | {c['ad']} | T:{c['t'] or '—'}, L:{c['l'] or '—'}, "
                f"Kredi:{c['kredi'] or '—'}, AKTS:{c['akts'] or '—'} | Ön şart: {c['on'] or 'yok'}"
                for c in courses
            ]
            body = (
                f"Endüstri Mühendisliği {mufredat_yili} Müfredatı — {havuz_label} "
                f"(toplam {len(courses)} ders):\n"
                + "\n".join(lines)
                + f"\n\nNot: Bu liste {mufredat_yili} müfredatına tabi öğrenciler için geçerlidir."
            )
            yield Chunk(
                id=f"ie_msec_havuz_ozet_{mufredat_yili}_{si}",
                text=body,
                metadata={
                    "tip": "secmeli_havuz_ozet",
                    "mufredat_yili": mufredat_yili,
                    "havuz": havuz_label,
                    "ders_sayisi": len(courses),
                    "kaynak": path.name,
                },
            )


def parse_ie_secmeli_csv(path: Path, mufredat_yili: str) -> Iterator[Chunk]:
    """Endüstri müfredatı seçmeli ders havuzu CSV: Ders Kodu,Ders Adı,Teo.,Lab.,Kredi,AKTS"""
    courses = []
    with path.open(encoding="utf-8") as f:
        reader = csv.DictReader(f)
        for ri, row in enumerate(reader, start=1):
            kod = (row.get("Ders Kodu") or "").strip()
            ad = (row.get("Ders Adı") or "").strip()
            t = (row.get("Teo.") or row.get("Teorik") or "").strip()
            l = (row.get("Lab.") or row.get("Lab") or "").strip()
            k = (row.get("Kredi") or "").strip()
            a = (row.get("AKTS") or "").strip()
            if not ad or not kod:
                continue
            courses.append({"kod": kod, "ad": ad, "t": t, "l": l, "kredi": k, "akts": a})

            text = (
                f"Endüstri Mühendisliği {mufredat_yili} Müfredatı — Seçmeli Ders Havuzu: "
                f"{kod} {ad}. Teorik: {t}, Lab: {l}, Kredi: {k}, AKTS: {a}. "
                f"Bu ders {mufredat_yili} müfredatı seçmeli olarak alınabilir."
            )
            yield Chunk(
                id=f"ie_sec_{mufredat_yili}_{kod.replace(' ','')}",
                text=text,
                metadata={
                    "tip": "secmeli_havuz",
                    "mufredat_yili": mufredat_yili,
                    "havuz": f"Endüstri Seçmeli Havuzu ({mufredat_yili})",
                    "ders_kodu": kod,
                    "ders_adi": ad,
                    "teorik": t,
                    "lab": l,
                    "kredi": k,
                    "akts": a,
                    "kaynak": path.name,
                },
            )

    if courses:
        lines = [
            f"- {c['kod']} | {c['ad']} | T:{c['t']}, L:{c['l']}, Kredi:{c['kredi']}, AKTS:{c['akts']}"
            for c in courses
        ]
        body = (
            f"Endüstri Mühendisliği {mufredat_yili} Müfredatı — Seçmeli Ders Havuzu "
            f"(toplam {len(courses)} ders):\n"
            + "\n".join(lines)
            + f"\n\nNot: Bu liste {mufredat_yili} müfredatına tabi öğrenciler için geçerlidir. "
              "Öğrenciler müfredatlarındaki seçmeli kredisi kadar bu havuzdan ders alır."
        )
        yield Chunk(
            id=f"ie_sec_havuz_ozet_{mufredat_yili}",
            text=body,
            metadata={
                "tip": "secmeli_havuz_ozet",
                "mufredat_yili": mufredat_yili,
                "havuz": f"Endüstri Seçmeli Havuzu ({mufredat_yili})",
                "ders_sayisi": len(courses),
                "kaynak": path.name,
            },
        )


# ----------------------------- Staj Kılavuzu (numaralı bölümler) -----------------------------

NUMBERED_SEC_RE = re.compile(r"^\s*(\d+)\.\s+([A-ZÇĞİÖŞÜ][^\n]{0,120})$")


def parse_ie_staj_kilavuzu(path: Path) -> Iterator[Chunk]:
    doc = Document(path)
    paragraphs = [p.text.strip() for p in doc.paragraphs]

    sec_idx = [i for i, t in enumerate(paragraphs) if NUMBERED_SEC_RE.match(t)]
    if not sec_idx:
        yield Chunk(
            id="ie_staj_kilavuz_full",
            text="IE Staj Kılavuzu:\n" + "\n".join(p for p in paragraphs if p),
            metadata={"tip": "staj_kilavuz", "dil": "tr", "kaynak": path.name},
        )
        return

    for k, start in enumerate(sec_idx):
        end = sec_idx[k + 1] if k + 1 < len(sec_idx) else len(paragraphs)
        m = NUMBERED_SEC_RE.match(paragraphs[start])
        sec_no, sec_title = m.group(1), m.group(2).strip()
        body = "\n".join(p for p in paragraphs[start + 1 : end] if p).strip()
        text = f"Endüstri Staj Kılavuzu — Bölüm {sec_no}. {sec_title}\n{body}"
        yield Chunk(
            id=f"ie_staj_sec_{sec_no}",
            text=text,
            metadata={
                "tip": "staj_kilavuz",
                "bolum_no": int(sec_no),
                "baslik": sec_title,
                "dil": "tr",
                "kaynak": path.name,
            },
        )


# ----------------------------- Syllabus (IE197/297/397) -----------------------------

def parse_ie_syllabus(path: Path) -> Iterator[Chunk]:
    """Syllabus docx'leri — tüm metin + tablolar düz metin olarak tek chunk."""
    doc = Document(path)
    parts = []
    # İsimden ders kodunu çıkar
    m = re.search(r"IE\s*(\d{3})", path.stem)
    ders_kodu = f"IE {m.group(1)}" if m else path.stem

    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            parts.append(t)
    for tbl in doc.tables:
        for row in tbl.rows:
            line = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
            if line:
                parts.append(line)

    body = "\n".join(parts)
    text = f"Syllabus — {ders_kodu} (Endüstri Stajı):\n{body}"
    yield Chunk(
        id=f"ie_syllabus_{ders_kodu.replace(' ','')}",
        text=text[:6000],
        metadata={
            "tip": "syllabus",
            "ders_kodu": ders_kodu,
            "dil": "en",
            "kaynak": path.name,
        },
    )


# ----------------------------- Main -----------------------------

IE_MUFREDAT_LEGACY = {}  # tüm yıllar artık temiz CSV'den
IE_MUFREDAT_NEW = {}     # 2025 docx artık devre dışı
# Temiz CSV müfredatları
# 2016: 2016-2020 girişliler  |  2021: 2021-2024 girişliler  |  2025: 2025+ girişliler
IE_MUFREDAT_CSV = {
    "2016": "ie_2016_2020_temiz.csv",
    "2021": "ie_2021_2024_temiz.csv",
    "2025": "ie_2025_temiz.csv",
}
# Tek-tablolu seçmeli havuz CSV'leri
IE_SECMELI_CSV = {
    "2016": "ie_2016_2020_secmeli.csv",
}
# Çift-tablolu seçmeli CSV'ler (havuz1, havuz2)
IE_DUAL_SECMELI_CSV = {
    "2021": ("ie_2021_2024_konsantrasyon.csv", "Konsantrasyon Alanı Seçmelisi", "Matematik Seçmelisi"),
}
# Çok-tablolu seçmeli CSV'ler (label satırlarıyla otomatik bölünür)
IE_MULTI_SECMELI_CSV = {
    "2025": ("ie_2025_secmeli.csv", "Bölüm Seçmelisi"),
}


def main():
    all_chunks: list[Chunk] = []

    for yil, fn in IE_MUFREDAT_LEGACY.items():
        p = RAW / fn
        if not p.exists():
            print(f"[!] bulunamadı: {p.name}")
            continue
        n = 0
        for ch in parse_ie_mufredat_legacy(p, yil):
            all_chunks.append(ch); n += 1
        print(f"[IE müfredat {yil}] {n} chunk")

    for yil, fn in IE_MUFREDAT_NEW.items():
        p = RAW / fn
        if not p.exists():
            print(f"[!] bulunamadı: {p.name}")
            continue
        n = 0
        for ch in parse_ie_mufredat_new(p, yil):
            all_chunks.append(ch); n += 1
        print(f"[IE müfredat {yil}] {n} chunk")

    for yil, fn in IE_MUFREDAT_CSV.items():
        p = RAW / fn
        if not p.exists():
            print(f"[!] bulunamadı: {p.name}")
            continue
        n = 0
        for ch in parse_ie_mufredat_csv(p, yil):
            all_chunks.append(ch); n += 1
        print(f"[IE müfredat {yil} CSV] {n} chunk")

    for yil, fn in IE_SECMELI_CSV.items():
        p = RAW / fn
        if not p.exists():
            print(f"[!] bulunamadı: {p.name}")
            continue
        n = 0
        for ch in parse_ie_secmeli_csv(p, yil):
            all_chunks.append(ch); n += 1
        print(f"[IE seçmeli {yil}] {n} chunk")

    for yil, (fn, h1, h2) in IE_DUAL_SECMELI_CSV.items():
        p = RAW / fn
        if not p.exists():
            print(f"[!] bulunamadı: {p.name}")
            continue
        n = 0
        for ch in parse_ie_dual_secmeli_csv(p, yil, h1, h2):
            all_chunks.append(ch); n += 1
        print(f"[IE dual seçmeli {yil}] {n} chunk")

    for yil, (fn, default_label) in IE_MULTI_SECMELI_CSV.items():
        p = RAW / fn
        if not p.exists():
            print(f"[!] bulunamadı: {p.name}")
            continue
        n = 0
        for ch in parse_ie_multi_secmeli_csv(p, yil, default_label):
            all_chunks.append(ch); n += 1
        print(f"[IE multi seçmeli {yil}] {n} chunk")

    p = RAW / "IE_STAJ_PROGRAMI_KILAVUZU (TR) - 29.09.2025.docx"
    if p.exists():
        n = 0
        for ch in parse_ie_staj_kilavuzu(p):
            all_chunks.append(ch); n += 1
        print(f"[IE staj kılavuz] {n} chunk")

    for fn in ["IE197_Syllabus.docx", "IE297_Syllabus.docx", "IE397_Syllabus.docx"]:
        p = RAW / fn
        if not p.exists():
            continue
        n = 0
        for ch in parse_ie_syllabus(p):
            all_chunks.append(ch); n += 1
        print(f"[IE syllabus {fn}] {n} chunk")

    out_path = OUT / "chunks_endustri.jsonl"
    with out_path.open("w", encoding="utf-8") as f:
        for ch in all_chunks:
            ch.metadata["bolum"] = BOLUM
            f.write(json.dumps(asdict(ch), ensure_ascii=False) + "\n")

    print(f"\n[OK] Toplam {len(all_chunks)} endüstri chunk -> {out_path}")


if __name__ == "__main__":
    main()
