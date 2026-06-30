"""
AGÜ Biyomühendislik (Bioengineering) bölümü ingest.
Kaynak: data/raw/AGU_Biyomuhendislik_Mufredat_2021.docx
  (Biyomühendislik Lisans yeterlilik veri tabanı, 17.05.2021)

Belge yapısı (python-docx tablo indeksleri sabit):
  - table[0]  : Program bilgileri (hakkında, hedefler, mezuniyet, kariyer ...)
  - table[2]  : Kredi/AKTS dağılım özeti
  - table[3..10]: 8 dönemlik ders programı (1.GÜZ ... 4.BAHAR)
  - table[11] : Alan seçmelileri (A: Biyomateryal, B: Genetik, C: Biyomedikal + diğer)
  - table[12] : GLB seçmelileri
  - table[13..]: Ders içerikleri (2 sütunlu: Kodu/İsim/AKTS/.../İçerik)

Çıktı: data/processed/chunks_biyomuhendislik.jsonl
bolum = "biyomuhendislik" | mufredat_yili = "2021"
"""
from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
RAW = ROOT / "data" / "raw"
OUT = ROOT / "data" / "processed"
OUT.mkdir(parents=True, exist_ok=True)

DOCX_PATH = RAW / "AGU_Biyomuhendislik_Mufredat_2021.docx"
BOLUM = "biyomuhendislik"
BOLUM_ADI = "Biyomühendislik"
MUFREDAT = "2021"
KAYNAK = DOCX_PATH.name

# Dönem ders programı tabloları: table index -> (donem, yil, sezon)
SEMESTER_TABLES = {
    3: (1, 1, "Güz"), 4: (2, 1, "Bahar"),
    5: (3, 2, "Güz"), 6: (4, 2, "Bahar"),
    7: (5, 3, "Güz"), 8: (6, 3, "Bahar"),
    9: (7, 4, "Güz"), 10: (8, 4, "Bahar"),
}


def _clean(s: str) -> str:
    return re.sub(r"\s+", " ", (s or "").replace("\n", " ")).strip()


def _tr_lower(s: str) -> str:
    """Türkçe-duyarlı küçük harf. Python'da 'İ'.lower() birleşik nokta (U+0307)
    üretip eşleşmeyi bozuyor; İ->i, I->ı yapıp combining dot'u temizler."""
    return (s or "").replace("İ", "i").replace("I", "ı").replace("̇", "").lower()


def _norm_int(s: str) -> str:
    """'3.0' -> '3', boşsa ''."""
    s = (s or "").strip()
    try:
        f = float(s.replace(",", "."))
        if f.is_integer():
            return str(int(f))
    except (ValueError, TypeError):
        pass
    return s


def _cell(row, i: int) -> str:
    cells = row.cells
    return _clean(cells[i].text) if i < len(cells) else ""


# ---------------- 1) Dönem ders programı + dönem özetleri ----------------

def parse_semesters(doc) -> list[dict]:
    chunks: list[dict] = []
    for ti, (donem, yil, sezon) in SEMESTER_TABLES.items():
        if ti >= len(doc.tables):
            continue
        table = doc.tables[ti]
        lines: list[str] = []
        for row in table.rows[1:]:  # ilk satır başlık
            kod = _cell(row, 0)
            ad = _cell(row, 1)
            if not ad or ad.lower().startswith("toplam") or ad.lower() == "ders adı":
                continue
            on = _cell(row, 2).strip("-").strip()
            teorik = _norm_int(_cell(row, 3))
            lab = _norm_int(_cell(row, 4))
            kredi = _norm_int(_cell(row, 5))
            akts = _norm_int(_cell(row, 6))
            if not kod:
                kod = "BENG-SEC"  # seçmeli placeholder (Teknik/Teknik olmayan Seçmeli)

            text = (
                f"Biyomühendislik {MUFREDAT} Müfredatı — {yil}. yıl {sezon} dönemi "
                f"({donem}. dönem): {kod} {ad}. "
                f"Teorik: {teorik or '—'}, Lab: {lab or '—'}, Kredi: {kredi or '—'}, "
                f"AKTS: {akts or '—'}. Ön şart: {on if on else 'yok'}."
            )
            chunks.append({
                "id": f"biyo_muf_{MUFREDAT}_d{donem}_{kod.replace(' ', '')}_{len(chunks)}",
                "text": text,
                "metadata": {
                    "bolum": BOLUM,
                    "tip": "mufredat",
                    "mufredat_yili": MUFREDAT,
                    "donem": donem,
                    "yil": yil,
                    "sezon": sezon,
                    "ders_kodu": kod,
                    "ders_adi": ad,
                    "on_sart": on,
                    "teorik": teorik,
                    "lab": lab,
                    "kredi": kredi,
                    "akts": akts,
                    "kaynak": KAYNAK,
                },
            })
            lines.append(
                f"- {kod} | {ad} | T:{teorik or '—'}, L:{lab or '—'}, "
                f"Kredi:{kredi or '—'}, AKTS:{akts or '—'} | Ön şart: {on or 'yok'}"
            )

        if lines:
            body = (
                f"Biyomühendislik {MUFREDAT} Müfredatı — {yil}. sınıf {sezon} yarıyılı "
                f"({donem}. dönem) — Toplam {len(lines)} ders:\n" + "\n".join(lines)
            )
            chunks.append({
                "id": f"biyo_muf_{MUFREDAT}_sem_summary_d{donem}",
                "text": body,
                "metadata": {
                    "bolum": BOLUM,
                    "tip": "donem_ozet",
                    "mufredat_yili": MUFREDAT,
                    "donem": donem,
                    "yil": yil,
                    "sezon": sezon,
                    "ders_sayisi": len(lines),
                    "kaynak": KAYNAK,
                },
            })
    return chunks


# ---------------- 2) Ders içerikleri (2 sütunlu tablolar) ----------------

# Anahtar normalizasyonu: belgede "İsmi/Adı", "Ön Şart/Özel Koşul" gibi varyasyonlar var
KEY_MAP = [
    ("kod", ["kodu", "kod"]),
    ("ad", ["ismi", "isim", "adı", "adi", "ad"]),
    ("saat", ["haftalık saati", "haftalik saati", "haftalık saat", "haftalik saat"]),
    ("kredi", ["kredisi", "kredi"]),
    ("akts", ["akts", "ects"]),
    ("seviye", ["seviye/yıl", "seviye / yıl", "seviye/yil", "seviye / yil", "seviye"]),
    ("donem_label", ["dönem", "donem"]),
    ("tip", ["tip"]),
    ("on_sart", ["ön şart(lar)", "ön şart", "on sart", "önşart", "onsart", "ön koşul"]),
    ("ozel_kosul", ["özel koşul(lar)", "ozel kosul", "özel koşul"]),
    ("koordinator", ["koordinatör(ler)", "kordinatör(ler)", "koordinatörler",
                     "kordinatörler", "koordinatör", "kordinatör"]),
    ("yer", ["yer"]),
    ("icerik", ["içerik", "icerik"]),
]


def _match_key(raw_key: str) -> str | None:
    rk = _tr_lower(_clean(raw_key)).rstrip(":")
    for canon, variants in KEY_MAP:
        if any(rk == v for v in variants):
            return canon
    for canon, variants in KEY_MAP:
        if any(rk.startswith(v) for v in variants):
            return canon
    return None


def _ders_id(kod: str) -> str:
    return re.sub(r"[^A-Za-z0-9]", "_", kod)


def parse_course_details(doc) -> list[dict]:
    chunks: list[dict] = []
    seen_codes: set[str] = set()
    for ti, table in enumerate(doc.tables):
        if ti < 13:  # ders içerik tabloları 13'ten başlıyor
            continue
        if len(table.columns) != 2:
            continue
        data: dict[str, str] = {}
        for row in table.rows:
            k = _match_key(_cell(row, 0))
            v = _cell(row, 1)
            if k and v and k not in data:
                data[k] = v
        kod = _clean(data.get("kod", ""))
        if not kod:
            continue
        # normalize "BENG205" -> "BENG 205"
        kod = re.sub(r"^([A-Za-zĞÜŞİÖÇ]+)\s*0*(\d+)$", lambda m: f"{m.group(1).upper()} {m.group(2)}", kod)
        if kod in seen_codes:
            continue
        seen_codes.add(kod)

        ad = data.get("ad", "")
        on = data.get("on_sart", "").strip("-").strip()
        ozel = data.get("ozel_kosul", "").strip("-").strip()
        # Bazı satırlarda ön şart "Özel Koşul" alanına yazılmış (ör. BENG 202)
        if not on and ozel:
            on = ozel
        akts = _norm_int(data.get("akts", ""))
        kredi = _norm_int(data.get("kredi", ""))
        icerik = data.get("icerik", "")
        koord = data.get("koordinator", "")

        text = (
            f"AGÜ {BOLUM_ADI} bölümü ders içeriği — {kod} {ad}. "
            f"Haftalık saat: {data.get('saat', '—')}. Kredi: {kredi or '—'}. AKTS: {akts or '—'}. "
            f"Seviye/Yıl: {data.get('seviye', '—')}. Dönem: {data.get('donem_label', '—')}. "
            f"Tip: {data.get('tip', '—')}. Ön şart: {on if on else 'yok'}. "
            + (f"Koordinatör: {koord}. " if koord else "")
            + (f"İçerik: {icerik}" if icerik else "")
        )
        chunks.append({
            "id": f"biyo_ders_{_ders_id(kod)}",
            "text": text[:5000],
            "metadata": {
                "bolum": BOLUM,
                "tip": "ders_icerik",
                "mufredat_yili": MUFREDAT,
                "ders_kodu": kod,
                "ders_adi": ad,
                "tip_ders": data.get("tip", ""),
                "donem_label": data.get("donem_label", ""),
                "kredi": kredi,
                "akts": akts,
                "on_sart": on if on else "yok",
                "koordinator": koord,
                "kaynak": KAYNAK,
            },
        })
    return chunks


# ---------------- 3) Alan seçmelileri (table 11) ----------------

def parse_electives(doc) -> list[dict]:
    if len(doc.tables) <= 11:
        return []
    table = doc.tables[11]
    chunks: list[dict] = []
    current_area = "Genel"
    area_courses: dict[str, list[str]] = {}
    for row in table.rows:
        c0 = _cell(row, 0)
        c1 = _cell(row, 1)
        # Başlık satırı (tüm hücreler aynı = alan adı)
        if c0 and c0 == c1:
            current_area = c0
            continue
        kod = c0
        ad = c1
        # Sütun başlığı satırlarını atla (Ders Kodu | Ders Adı ...)
        if not kod or not ad or _tr_lower(kod) in ("ders kodu", "kodu", "kod"):
            continue
        # son sütunlar kredi/akts
        cells = [_clean(c.text) for c in row.cells]
        kredi = _norm_int(cells[-2]) if len(cells) >= 2 else ""
        akts = _norm_int(cells[-1]) if len(cells) >= 1 else ""
        area_courses.setdefault(current_area, []).append(
            f"- {kod} | {ad} | Kredi: {kredi or '—'}, AKTS: {akts or '—'}"
        )

    for area, lines in area_courses.items():
        body = (
            f"AGÜ {BOLUM_ADI} bölümü — Alan Seçmelileri ({area}), toplam {len(lines)} ders:\n"
            + "\n".join(lines)
        )
        slug = re.sub(r"[^a-z0-9]+", "_", area.lower()).strip("_")
        chunks.append({
            "id": f"biyo_secmeli_{slug}",
            "text": body[:5000],
            "metadata": {
                "bolum": BOLUM,
                "tip": "secmeli_listesi",
                "mufredat_yili": MUFREDAT,
                "kategori": area,
                "ders_sayisi": len(lines),
                "kaynak": KAYNAK,
            },
        })
    return chunks


# ---------------- 4) Program genel bilgi (table 0) ----------------

def parse_program_info(doc) -> list[dict]:
    if not doc.tables:
        return []
    table = doc.tables[0]
    info: dict[str, str] = {}
    for row in table.rows:
        k = _clean(_cell(row, 0))
        v = _clean(_cell(row, 1))
        if k and v and k not in info:
            info[k] = v

    # İlgilenilen alanlar
    parts = []
    for label in [
        "Program Hakkında", "Program Hedefleri", "Kazanılan Derece",
        "Öğrenim Süresi ve Kredisi", "Eğitim Türü", "Kabul Koşulları",
        "Mezuniyet Koşulları ve Kuralları", "Mezunların Mesleki Profili ve İstihdam Olanakları",
        "Üst Derece Programlarına Geçiş",
    ]:
        if info.get(label):
            parts.append(f"{label}: {info[label]}")
    text = f"AGÜ {BOLUM_ADI} Lisans Programı — Genel Bilgiler.\n\n" + "\n\n".join(parts)

    chunks = [{
        "id": "biyo_program_genel",
        "text": text[:5000],
        "metadata": {
            "bolum": BOLUM,
            "tip": "program_genel",
            "mufredat_yili": MUFREDAT,
            "kaynak": KAYNAK,
        },
    }]

    # Mezuniyet kuralları ayrı chunk (sık sorulur)
    if info.get("Mezuniyet Koşulları ve Kuralları"):
        chunks.append({
            "id": "biyo_mezuniyet_kurallari",
            "text": (
                f"AGÜ {BOLUM_ADI} bölümü mezuniyet koşulları ve kuralları:\n"
                f"{info['Mezuniyet Koşulları ve Kuralları']}\n\n"
                f"Öğrenim süresi/kredi: {info.get('Öğrenim Süresi ve Kredisi', '4 yıl, 240 AKTS')}"
            )[:5000],
            "metadata": {
                "bolum": BOLUM,
                "tip": "mezuniyet_kurallari",
                "mufredat_yili": MUFREDAT,
                "kaynak": KAYNAK,
            },
        })
    return chunks


def make_alan_secimi() -> dict:
    text = (
        "AGÜ Biyomühendislik bölümü — 4. Sınıf Alan Seçimi (Konsantrasyon/Track Seçimi):\n\n"
        "Öğrenciler 4. sınıfa geldiklerinde aşağıdaki üç alandan birini seçebilir:\n"
        "- A) Biyomateryal ve Doku Mühendisliği (Biomaterials & Tissue Engineering)\n"
        "- B) Genetik ve Biyoproses Mühendisliği (Genetics & Bioprocessing)\n"
        "- C) Biyomedikal Elektroniği (Biomedical Electronics)\n\n"
        "KURALLAR:\n"
        "- Seçilen alandan mezun olmak için Alan Seçmeli (CA Elective / Concentration Area Elective) "
        "derslerinin EN AZ YARISININ seçilen alandan olması gerekir.\n"
        "- Ayrıca Capstone Project (Bitirme Projesi BENG 491/492) derslerinin de bu alandan olması gerekir.\n"
        "- Öğrenciler hiç alan seçmeden de devam edebilir (alan seçimi zorunlu değildir).\n\n"
        "Soru: 'Biyomühendislikte alan seçimi nasıl yapılır?' / 'Konsantrasyon seçmek zorunda mıyım?' / "
        "'Hangi alanları seçebilirim?' / 'Alan seçmeden mezun olabilir miyim?' → Yukarıdaki kurallar geçerlidir. "
        "Alan seçimi zorunlu değildir; seçilirse CA Elective derslerinin en az yarısı ve Capstone o alandan olmalıdır."
    )
    return {
        "id": "biyo_alan_secimi",
        "text": text,
        "metadata": {
            "bolum": BOLUM,
            "tip": "alan_secimi",
            "mufredat_yili": MUFREDAT,
            "kaynak": "bolum_bilgilendirme",
        },
    }


def make_overview() -> dict:
    text = (
        "AGÜ Biyomühendislik (Bioengineering, ders kodu: BENG) lisans programı. "
        "Abdullah Gül Üniversitesi Mühendislik Fakültesi altında, 4 yıllık (8 dönem) lisans programı, "
        "toplam 240 AKTS (bir yıl İngilizce Hazırlık hariç). Tek aktif müfredat: 2021. "
        "Biyomühendislik; tıp ve temel bilimlerin ilkelerini malzeme ve mühendislik bilimi ile "
        "birleştiren disiplinlerarası bir alandır; biyomedikal hesaplama ve görüntüleme, biyomedikal "
        "cihaz teknolojisi, hücre ve moleküler mühendislik, rejeneratif tıp gibi konulara odaklanır. "
        "Program 3. sınıfta alan seçimi içerir: A) Biyomateryal ve Doku Mühendisliği, "
        "B) Genetik ve Biyoproses Mühendisliği, C) Biyomedikal Elektroniği. "
        "Zorunlu dersler BENG kodludur (ör. BENG 101 Biyomühendisliğe Giriş, BENG 201 Biyokimya, "
        "BENG 491/492 Bitirme Projesi I-II, BENG 493 Yaz Stajı). "
        "Mezuniyet için minimum genel not ortalaması 2,00."
    )
    return {
        "id": "biyo_program_overview",
        "text": text,
        "metadata": {
            "bolum": BOLUM,
            "tip": "program_bilgi",
            "mufredat_yili": MUFREDAT,
            "kaynak": "program_tanitim",
        },
    }


def main():
    doc = Document(DOCX_PATH)
    chunks: list[dict] = []
    chunks.append(make_overview())
    chunks.append(make_alan_secimi())
    chunks.extend(parse_program_info(doc))
    chunks.extend(parse_semesters(doc))
    chunks.extend(parse_course_details(doc))
    chunks.extend(parse_electives(doc))

    out_path = OUT / "chunks_biyomuhendislik.jsonl"
    with out_path.open("w", encoding="utf-8") as f:
        for ch in chunks:
            f.write(json.dumps(ch, ensure_ascii=False) + "\n")

    print(f"[OK] {len(chunks)} biyomühendislik chunk -> {out_path}")
    by_tip: dict[str, int] = {}
    for c in chunks:
        t = c["metadata"]["tip"]
        by_tip[t] = by_tip.get(t, 0) + 1
    for t, n in sorted(by_tip.items()):
        print(f"     - {t}: {n}")


if __name__ == "__main__":
    main()
