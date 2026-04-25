"""
AGÜ Bilgisayar Mühendisliği RAG Asistanı — Ingest Pipeline
Tüm ham dokümanları data/raw/ içinden okur,
yapılandırılmış chunk'lara dönüştürüp data/processed/chunks.jsonl yazar.
"""
from __future__ import annotations

import csv
import json
import re
from dataclasses import dataclass, field, asdict
from pathlib import Path
from typing import Iterator

import pdfplumber
from docx import Document
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
RAW = ROOT / "data" / "raw"
OUT = ROOT / "data" / "processed"
OUT.mkdir(parents=True, exist_ok=True)


@dataclass
class Chunk:
    id: str
    text: str
    metadata: dict = field(default_factory=dict)


# ----------------------------- Müfredat (tablolu docx) -----------------------------

MUFREDAT_FILES = {
    # 2016, 2021, 2023: temiz CSV ile yapılıyor (parse_mufredat_csv)
    "2025": "CMP_Liste_2025-1 (2).docx",
}

# Temiz CSV müfredatları
# 2016: 2016-2020 girişliler  |  2021: 2021-2022 girişliler  |  2023: 2023-2024 girişliler
MUFREDAT_CSV = {
    "2016": "comp_2016_2021_temiz.csv",
    "2021": "comp_2021_2022_temiz.csv",
    "2023": "comp_2023_2024_temiz.csv",
}

DONEM_RE = re.compile(r"\(Dönem\s*(\d+)\)\s*-\s*Yıl\s*(\d+),\s*(\w+)\s*YARIYILI", re.IGNORECASE)
POOL_RE = re.compile(r"(seçmeli|secmeli|havuz)", re.IGNORECASE)


def _col_map(header_cells: list[str]) -> dict:
    """Header ismi -> kolon indexi. TR/EN varyantlarını tanır."""
    norm = {}
    aliases = {
        "ders_kodu": ["ders kodu", "kod", "code"],
        "ders_adi": ["ders adı", "ders adi", "ad", "name"],
        "on_sart": ["ön şart", "on sart", "ön sart", "prereq", "prerequisite"],
        "teorik": ["teorik", "teo.", "teo", "theoretical", "theory"],
        "lab": ["laboratuvar", "lab", "practical"],
        "kredi": ["kredi", "credit"],
        "akts": ["akts", "ects"],
    }
    for i, h in enumerate(header_cells):
        hl = h.strip().lower()
        for key, alts in aliases.items():
            if any(a == hl or hl.startswith(a) for a in alts):
                norm.setdefault(key, i)
                break
    return norm


def _cell(cells: list[str], cmap: dict, key: str) -> str:
    i = cmap.get(key)
    if i is None or i >= len(cells):
        return ""
    return cells[i].strip()


def _blankish(v: str) -> bool:
    return v is None or str(v).strip() in ("", "-", "—", "–", "N/A", "n/a")


def _synth_code(ders_adi: str) -> str:
    ad = (ders_adi or "").lower()
    if "bilim" in ad:
        return "COMP-SCI-ELEC"
    if "matematik" in ad or "math" in ad:
        return "COMP-MATH-ELEC"
    if "teknik" in ad or "technical" in ad:
        return "COMP-TECH-ELEC"
    if "non-technical" in ad or "non technical" in ad:
        return "COMP-NT-ELEC"
    if "küresel" in ad or "global" in ad:
        return "GLB-ELEC"
    if "türk" in ad or "turk" in ad:
        return "TURK-ELEC"
    if "tarih" in ad or "history" in ad:
        return "HIST-ELEC"
    return "COMP-ELEC"


def _fmt_hours(teorik, lab, kredi, akts, kod="", ders_adi="") -> str:
    t_b, l_b, k_b, a_b = _blankish(teorik), _blankish(lab), _blankish(kredi), _blankish(akts)
    ad_low = (ders_adi or "").lower()
    kod_up = (kod or "").upper().replace(" ", "")
    is_internship = (
        "summer training" in ad_low or "staj" in ad_low
        or kod_up in ("COMP197", "COMP297", "COMP397", "COMP497")
    )
    is_elective_placeholder = (
        not kod or kod.startswith("COMP-") and "ELEC" in kod.upper()
        or "ELEC" in kod_up or kod_up.startswith("MATHEX") or kod_up.startswith("SCIX")
        or "seçmeli" in ad_low or "secmeli" in ad_low or "elective" in ad_low
    )

    def _val(v, blank):
        if not blank:
            return str(v).strip()
        if is_internship:
            return "—"
        if is_elective_placeholder:
            return "seçim sonrası belirlenir"
        return "belirtilmemiş"

    base = f"Teorik: {_val(teorik,t_b)}, Lab: {_val(lab,l_b)}, Kredi: {_val(kredi,k_b)}, AKTS: {_val(akts,a_b)}."
    if is_internship and (t_b or l_b or k_b):
        base += " (Staj dersi — sınıf içi ders saati yoktur.)"
    elif is_elective_placeholder and (t_b or k_b):
        base += " (Seçmeli havuzdan seçilen dersin saat/kredi bilgileri o derse göre değişir.)"
    return base


def _iter_body(doc: Document) -> Iterator[tuple[str, object]]:
    """Body içinde paragraf ve tabloları sırayla dolaş."""
    p_idx, t_idx = 0, 0
    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:p"):
            yield "p", doc.paragraphs[p_idx]
            p_idx += 1
        elif child.tag == qn("w:tbl"):
            yield "t", doc.tables[t_idx]
            t_idx += 1


def parse_mufredat(path: Path, mufredat_yili: str) -> Iterator[Chunk]:
    doc = Document(path)
    # "context": bir sonraki tabloyu tanımlayan son anlamlı paragraf
    cur_donem = cur_yil = cur_sezon = None
    cur_pool = None  # ör. "Teknik Seçmeli", "GLB1XX Seçmeli"
    # Dönem özeti için: her dönemdeki tüm dersleri biriktir
    semester_courses: dict[tuple, list[str]] = {}
    semester_meta: dict[tuple, dict] = {}

    for kind, el in _iter_body(doc):
        if kind == "p":
            txt = el.text.strip()
            if not txt:
                continue
            m = DONEM_RE.search(txt)
            if m:
                cur_donem, cur_yil, cur_sezon = m.group(1), m.group(2), m.group(3)
                cur_pool = None
            elif POOL_RE.search(txt) and len(txt) < 80:
                cur_pool = txt
                cur_donem = None  # havuz tablosu geldi, dönem sıfırla
        elif kind == "t":
            rows = el.rows
            if not rows:
                continue
            header = [c.text.strip() for c in rows[0].cells]
            cmap = _col_map(header)
            if "ders_kodu" not in cmap or "ders_adi" not in cmap:
                continue  # anlaşılmayan tablo

            for ri, row in enumerate(rows[1:], start=1):
                cells = [c.text.strip() for c in row.cells]
                if not any(cells):
                    continue
                ders_kodu = _cell(cells, cmap, "ders_kodu")
                ders_adi = _cell(cells, cmap, "ders_adi")
                on_sart = _cell(cells, cmap, "on_sart")
                teorik = _cell(cells, cmap, "teorik")
                lab = _cell(cells, cmap, "lab")
                kredi = _cell(cells, cmap, "kredi")
                akts = _cell(cells, cmap, "akts")

                if not ders_adi or "TOPLAM" in ders_adi.upper() or ders_adi.upper() == "TOTAL":
                    continue
                # Dipnot/junk satırları (örn "(*) Bilim Seçmeli: ...") atla
                if ders_adi.startswith("(*") or ders_adi == ders_kodu:
                    continue

                kod_effective = ders_kodu if ders_kodu else _synth_code(ders_adi)
                hours_str = _fmt_hours(teorik, lab, kredi, akts, kod_effective, ders_adi)

                if cur_donem:
                    text = (
                        f"{mufredat_yili} müfredatı — {cur_yil}. yıl {cur_sezon} dönemi "
                        f"({cur_donem}. dönem): {kod_effective} {ders_adi}. "
                        f"{hours_str} "
                        f"Ön şart: {on_sart if on_sart else 'yok'}."
                    )
                    # Dönem özetine ekle
                    key = (int(cur_donem),)
                    short_line = (
                        f"- {kod_effective} | {ders_adi} | "
                        f"T:{teorik or '—'}, L:{lab or '—'}, Kredi:{kredi or '—'}, AKTS:{akts or '—'} | "
                        f"Ön şart: {on_sart or 'yok'}"
                    )
                    semester_courses.setdefault(key, []).append(short_line)
                    semester_meta.setdefault(key, {"yil": int(cur_yil), "sezon": cur_sezon})
                    md = {
                        "tip": "mufredat",
                        "mufredat_yili": mufredat_yili,
                        "donem": int(cur_donem),
                        "yil": int(cur_yil),
                        "sezon": cur_sezon,
                        "ders_kodu": kod_effective,
                        "ders_adi": ders_adi,
                        "on_sart": on_sart,
                        "teorik": teorik,
                        "lab": lab,
                        "kredi": kredi,
                        "akts": akts,
                        "kaynak": path.name,
                    }
                    cid = f"muf_{mufredat_yili}_d{cur_donem}_r{ri}_{kod_effective.replace(' ', '')}"
                else:
                    pool = cur_pool or "Seçmeli Havuz"
                    text = (
                        f"{mufredat_yili} müfredatı — {pool} havuzu: {kod_effective} {ders_adi}. "
                        f"{hours_str} "
                        f"Ön şart: {on_sart if on_sart else 'yok'}."
                    )
                    md = {
                        "tip": "secmeli_havuz",
                        "mufredat_yili": mufredat_yili,
                        "havuz": pool,
                        "ders_kodu": kod_effective,
                        "ders_adi": ders_adi,
                        "on_sart": on_sart,
                        "teorik": teorik,
                        "lab": lab,
                        "kredi": kredi,
                        "akts": akts,
                        "kaynak": path.name,
                    }
                    cid = f"muf_{mufredat_yili}_pool_{pool[:12].replace(' ','_')}_r{ri}_{kod_effective.replace(' ','')}"
                yield Chunk(id=cid, text=text, metadata=md)

    # Dönem başına özet chunk (makine tarzı — LLM'in eksiksiz kopyalaması için)
    for (donem,), lines in sorted(semester_courses.items()):
        meta = semester_meta[(donem,)]
        yil_g = meta["yil"]
        sezon = meta["sezon"]
        body = (
            f"Bilgisayar Mühendisliği {mufredat_yili} Müfredatı — "
            f"{yil_g}. sınıf {sezon} yarıyılı ({donem}. dönem) — Toplam {len(lines)} ders:\n"
            + "\n".join(lines)
        )
        yield Chunk(
            id=f"muf_{mufredat_yili}_sem_summary_d{donem}",
            text=body,
            metadata={
                "tip": "donem_ozet",
                "mufredat_yili": mufredat_yili,
                "donem": donem,
                "yil": yil_g,
                "sezon": sezon,
                "ders_sayisi": len(lines),
                "kaynak": path.name,
            },
        )


# ----------------------------- Müfredat: Temiz CSV (2016/2021) -----------------------------

SEZON_BY_DONEM = {1: "Güz", 2: "Bahar", 3: "Güz", 4: "Bahar", 5: "Güz", 6: "Bahar", 7: "Güz", 8: "Bahar"}


def parse_mufredat_csv(path: Path, mufredat_yili: str) -> Iterator[Chunk]:
    """Kullanıcının elle temizlediği CSV: Dönem,Ders Kodu,Ders Adı,Ön Şart,Teorik,Lab.,Kredi,AKTS"""
    semester_courses: dict[int, list[str]] = {}
    semester_meta: dict[int, dict] = {}

    with path.open(encoding="utf-8") as f:
        reader = csv.DictReader(f)
        for ri, row in enumerate(reader, start=1):
            try:
                donem = int(str(row.get("Dönem", "")).strip())
            except (ValueError, TypeError):
                continue
            ders_kodu = (row.get("Ders Kodu") or "").strip()
            ders_adi = (row.get("Ders Adı") or "").strip()
            on_sart = (row.get("Ön Şart") or "").strip().strip("-").strip()
            teorik = (row.get("Teorik") or "").strip()
            lab = (row.get("Lab.") or row.get("Lab") or "").strip()
            kredi = (row.get("Kredi") or "").strip()
            akts = (row.get("AKTS") or "").strip()

            if not ders_adi:
                continue
            if ders_kodu in ("-", ""):
                ders_kodu = _synth_code(ders_adi)

            yil = (donem - 1) // 2 + 1
            sezon = SEZON_BY_DONEM.get(donem, "")
            hours_str = _fmt_hours(teorik, lab, kredi, akts, ders_kodu, ders_adi)

            text = (
                f"{mufredat_yili} müfredatı — {yil}. yıl {sezon} dönemi "
                f"({donem}. dönem): {ders_kodu} {ders_adi}. "
                f"{hours_str} "
                f"Ön şart: {on_sart if on_sart else 'yok'}."
            )
            md = {
                "tip": "mufredat",
                "mufredat_yili": mufredat_yili,
                "donem": donem,
                "yil": yil,
                "sezon": sezon,
                "ders_kodu": ders_kodu,
                "ders_adi": ders_adi,
                "on_sart": on_sart,
                "teorik": teorik,
                "lab": lab,
                "kredi": kredi,
                "akts": akts,
                "kaynak": path.name,
            }
            yield Chunk(
                id=f"muf_{mufredat_yili}_csv_d{donem}_r{ri}_{ders_kodu.replace(' ','')}",
                text=text,
                metadata=md,
            )

            short_line = (
                f"- {ders_kodu} | {ders_adi} | "
                f"T:{teorik or '—'}, L:{lab or '—'}, Kredi:{kredi or '—'}, AKTS:{akts or '—'} | "
                f"Ön şart: {on_sart or 'yok'}"
            )
            semester_courses.setdefault(donem, []).append(short_line)
            semester_meta.setdefault(donem, {"yil": yil, "sezon": sezon})

    # Dönem özeti
    for donem, lines in sorted(semester_courses.items()):
        meta = semester_meta[donem]
        body = (
            f"Bilgisayar Mühendisliği {mufredat_yili} Müfredatı — "
            f"{meta['yil']}. sınıf {meta['sezon']} yarıyılı ({donem}. dönem) — Toplam {len(lines)} ders:\n"
            + "\n".join(lines)
        )
        yield Chunk(
            id=f"muf_{mufredat_yili}_csv_sem_summary_d{donem}",
            text=body,
            metadata={
                "tip": "donem_ozet",
                "mufredat_yili": mufredat_yili,
                "donem": donem,
                "yil": meta["yil"],
                "sezon": meta["sezon"],
                "ders_sayisi": len(lines),
                "kaynak": path.name,
            },
        )


# ----------------------------- Teknik Seçmeli Havuzu (CSV) -----------------------------

def parse_teknik_secmeli_csv(path: Path) -> Iterator[Chunk]:
    """Bilgisayar Müh. teknik seçmeli ders havuzu — tüm müfredat yıllarında geçerli."""
    courses = []
    with path.open(encoding="utf-8") as f:
        reader = csv.DictReader(f)
        for ri, row in enumerate(reader, start=1):
            kod = (row.get("Ders Kodu") or "").strip()
            ad = (row.get("Ders Adı") or "").strip()
            t = (row.get("Teo.") or row.get("Teorik") or "").strip()
            l = (row.get("Lab") or "").strip()
            k = (row.get("Kredi") or "").strip()
            a = (row.get("AKTS") or "").strip()
            if not ad or not kod:
                continue
            courses.append({"kod": kod, "ad": ad, "t": t, "l": l, "kredi": k, "akts": a})

            # Per-ders chunk (tek ders sorgularına bulunabilir olsun)
            text = (
                f"Bilgisayar Mühendisliği Teknik Seçmeli Havuzu — {kod} {ad}. "
                f"Teorik: {t}, Lab: {l}, Kredi: {k}, AKTS: {a}. "
                f"Bu ders teknik seçmeli olarak alınabilir (tüm müfredat yıllarında geçerli)."
            )
            yield Chunk(
                id=f"tech_elec_{kod.replace(' ','')}",
                text=text,
                metadata={
                    "tip": "teknik_secmeli",
                    "havuz": "Teknik Seçmeli (Bilgisayar Müh.)",
                    "ders_kodu": kod,
                    "ders_adi": ad,
                    "teorik": t,
                    "lab": l,
                    "kredi": k,
                    "akts": a,
                    "kaynak": path.name,
                },
            )

    # Topluca özet chunk — "teknik seçmelileri listele" gibi sorgular için
    if courses:
        lines = [
            f"- {c['kod']} | {c['ad']} | T:{c['t']}, L:{c['l']}, Kredi:{c['kredi']}, AKTS:{c['akts']}"
            for c in courses
        ]
        body = (
            f"Bilgisayar Mühendisliği — Teknik Seçmeli Ders Havuzu (toplam {len(courses)} ders, "
            f"tüm müfredat yıllarında geçerli):\n"
            + "\n".join(lines)
            + "\n\nNot: Öğrenciler bu havuzdan müfredatlarındaki teknik seçmeli kredisi kadar ders alır."
        )
        yield Chunk(
            id="tech_elec_havuz_ozet",
            text=body,
            metadata={
                "tip": "teknik_secmeli_ozet",
                "havuz": "Teknik Seçmeli (Bilgisayar Müh.)",
                "ders_sayisi": len(courses),
                "kaynak": path.name,
            },
        )


# ----------------------------- Staj Yönergesi (MADDE) -----------------------------

MADDE_RE = re.compile(r"^MADDE\s*(\d+)\s*[-–]?\s*(.*)$")


def parse_yonerge(path: Path) -> Iterator[Chunk]:
    doc = Document(path)
    paragraphs = [p.text.strip() for p in doc.paragraphs]

    # Her MADDE için: önceki non-empty non-MADDE satırı "başlık" (Amaç, Kapsam, ...)
    # İçerik: MADDE satırından sonraki non-empty satırlar, sıradaki MADDE'ye kadar.
    madde_positions = []
    for i, t in enumerate(paragraphs):
        if MADDE_RE.match(t):
            madde_positions.append(i)

    for idx, start in enumerate(madde_positions):
        end = madde_positions[idx + 1] if idx + 1 < len(madde_positions) else len(paragraphs)
        m = MADDE_RE.match(paragraphs[start])
        madde_no = m.group(1)
        first_line = m.group(2).strip()

        # Başlık: start'tan yukarıda en yakın kısa, MADDE olmayan satır
        baslik = ""
        for j in range(start - 1, max(-1, start - 5), -1):
            s = paragraphs[j]
            if s and not MADDE_RE.match(s) and len(s) < 60:
                baslik = s
                break

        body_lines = [first_line] if first_line else []
        for j in range(start + 1, end):
            s = paragraphs[j]
            if s and not MADDE_RE.match(s):
                body_lines.append(s)
        body = "\n".join(body_lines).strip()

        text = f"Staj Yönergesi — MADDE {madde_no} ({baslik}):\n{body}" if baslik else f"Staj Yönergesi — MADDE {madde_no}:\n{body}"
        yield Chunk(
            id=f"yon_madde_{madde_no}",
            text=text,
            metadata={
                "tip": "yonerge",
                "madde_no": int(madde_no),
                "baslik": baslik,
                "dil": "tr",
                "kaynak": path.name,
            },
        )


# ----------------------------- Internship Handbook (EN düz metin) -----------------------------

HANDBOOK_SECTION_RE = re.compile(r"^\s*(\d+)\.\s+([A-Z][^\n]{0,80})$")


def parse_handbook(path: Path) -> Iterator[Chunk]:
    doc = Document(path)
    paragraphs = [p.text.strip() for p in doc.paragraphs]

    # Numaralı başlıklara göre bölümle
    section_idx = []
    for i, t in enumerate(paragraphs):
        if HANDBOOK_SECTION_RE.match(t):
            section_idx.append(i)

    if not section_idx:
        # Fallback: tek chunk
        yield Chunk(
            id="handbook_full",
            text="Internship Handbook:\n" + "\n".join(p for p in paragraphs if p),
            metadata={"tip": "handbook", "dil": "en", "kaynak": path.name},
        )
        return

    for k, start in enumerate(section_idx):
        end = section_idx[k + 1] if k + 1 < len(section_idx) else len(paragraphs)
        m = HANDBOOK_SECTION_RE.match(paragraphs[start])
        sec_no, sec_title = m.group(1), m.group(2).strip()
        body = "\n".join(p for p in paragraphs[start + 1 : end] if p).strip()
        text = f"Internship Handbook — Section {sec_no}. {sec_title}\n{body}"
        yield Chunk(
            id=f"hb_sec_{sec_no}",
            text=text,
            metadata={
                "tip": "handbook",
                "bolum_no": int(sec_no),
                "baslik": sec_title,
                "dil": "en",
                "kaynak": path.name,
            },
        )


# ----------------------------- Ders Katalog PDF -----------------------------

COURSE_RECORD_MARKER = "COURSE RECORD"
CODE_RE = re.compile(r"^Code\s+([A-Z]{2,5}\s*\d{3})", re.MULTILINE)
NAME_RE = re.compile(r"^Name\s+(.+)$", re.MULTILINE)
SEMESTER_RE = re.compile(r"^Semester/Year\s+(.+)$", re.MULTILINE)
PREREQ_RE = re.compile(r"^Prerequisites\s+(.+)$", re.MULTILINE)


def parse_katalog(path: Path) -> Iterator[Chunk]:
    with pdfplumber.open(path) as pdf:
        full = "\n".join((page.extract_text() or "") for page in pdf.pages)

    # COURSE RECORD ile böl; ilk parça intro/özet olabilir
    parts = full.split(COURSE_RECORD_MARKER)
    for i, part in enumerate(parts[1:], start=1):
        # Sonraki bloğun başına kadar olan kısım bu dersin
        body = part.strip()
        code_m = CODE_RE.search(body)
        name_m = NAME_RE.search(body)
        sem_m = SEMESTER_RE.search(body)
        pre_m = PREREQ_RE.search(body)

        if not code_m:
            continue
        ders_kodu = code_m.group(1).strip()
        ders_adi = name_m.group(1).strip() if name_m else ""
        sezon_yil = sem_m.group(1).strip() if sem_m else ""
        on_sart = pre_m.group(1).strip() if pre_m else ""

        text = f"Ders Kataloğu — {ders_kodu} {ders_adi}\nSemester: {sezon_yil}\nPrerequisites: {on_sart}\n\n{body}"
        yield Chunk(
            id=f"kat_{ders_kodu.replace(' ', '')}",
            text=text[:4000],  # çok uzun olursa kırp
            metadata={
                "tip": "ders_katalog",
                "ders_kodu": ders_kodu,
                "ders_adi": ders_adi,
                "sezon_yil": sezon_yil,
                "on_sart": on_sart,
                "dil": "en",
                "kaynak": path.name,
            },
        )


# ----------------------------- Main -----------------------------

def main():
    all_chunks: list[Chunk] = []

    # Müfredatlar (docx)
    for yil, fn in MUFREDAT_FILES.items():
        p = RAW / fn
        if not p.exists():
            print(f"[!] bulunamadı: {p.name}")
            continue
        n = 0
        for ch in parse_mufredat(p, yil):
            all_chunks.append(ch)
            n += 1
        print(f"[müfredat {yil}] {n} chunk")

    # Müfredatlar (temiz CSV — 2016/2021)
    for yil, fn in MUFREDAT_CSV.items():
        p = RAW / fn
        if not p.exists():
            print(f"[!] bulunamadı: {p.name}")
            continue
        n = 0
        for ch in parse_mufredat_csv(p, yil):
            all_chunks.append(ch)
            n += 1
        print(f"[müfredat {yil} CSV] {n} chunk")

    # Teknik Seçmeli Havuzu
    p = RAW / "comp_teknik_secmeli.csv"
    if p.exists():
        n = 0
        for ch in parse_teknik_secmeli_csv(p):
            all_chunks.append(ch); n += 1
        print(f"[teknik seçmeli] {n} chunk")

    # Yönerge
    p = RAW / "AGU_Bilgisayar_Staj_Yonergesi_2026.docx"
    if p.exists():
        n = 0
        for ch in parse_yonerge(p):
            all_chunks.append(ch); n += 1
        print(f"[yönerge] {n} chunk")

    # Handbook
    p = RAW / "AGU-COMP-Internship-Handbook.docx"
    if p.exists():
        n = 0
        for ch in parse_handbook(p):
            all_chunks.append(ch); n += 1
        print(f"[handbook] {n} chunk")

    # Katalog
    p = RAW / "DERS_KATALOG_COMP_v1.pdf"
    if p.exists():
        n = 0
        for ch in parse_katalog(p):
            all_chunks.append(ch); n += 1
        print(f"[katalog] {n} chunk")

    out_path = OUT / "chunks.jsonl"
    with out_path.open("w", encoding="utf-8") as f:
        for ch in all_chunks:
            ch.metadata["bolum"] = "bilgisayar"
            f.write(json.dumps(asdict(ch), ensure_ascii=False) + "\n")

    print(f"\n[OK] Toplam {len(all_chunks)} chunk -> {out_path}")


if __name__ == "__main__":
    main()
